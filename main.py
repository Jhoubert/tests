import os
import sys
import re
from docx import Document

base_path = "./documents"

#doc.tables[0]._cells[1].text

def read_test(doc):
    questions=[]
    question = {
        "question":"",
        "correct_answer": "",
        "answers": [],
        "ok": False
    }
    n = -1
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) == 2:
                print("This is a student document")
                return None 
            """
            cell_0 # Should be empty or right answer
            cell_1 # Should be the question number or letter
            cell_2 # Should be the text :)
            """
            
            if re.match("[0-9]+[\.\-\s]{1,3}", cells[1].text.strip()):
                n+=1
                questions.append(question.copy())
                questions[n]["question"] = cells[2].text.strip()
            elif re.match("[a-zA-Z]\)", cells[1].text.strip()):
                questions[n]["answers"].append(cells[2].text.strip())
                if len(cells[0].text.strip()) > 0:
                    questions[n]["ok"] = True
                    questions[n]["correct_answer"] = cells[2].text.strip()
                else:
                    questions[n]["answers"].append(cells[2].text.strip())
            else:
                questions[n]["ok"] = False
                print("unrecognized row format")
        
    return questions

search = sys.argv[1]
found = set()
documents = []
for root, d_names, d_files in os.walk(base_path):
    for file in d_files:
        if file.endswith(".docx"):
            file_path = root + "/" + file
            print("\nOpening:", file_path)
            doc = Document(file_path)
            questions = read_test(doc)
            if questions:
                print("Encontradas: "+str(len(questions) if questions else 0)+" preguntas")
                for q in questions:
                    print(" - - - - - - - - - - ")
                    print("Q:", q["question"])
                    print("A:", q["correct_answer"])
                    if search in q["question"] or search in q["correct_answer"] or search in "\n".join(q["answers"]):
                        found.add(file_path)
                documents.append(questions)

            print("-------------------------------------------")
            print("-------------------------------------------")
    
    print("\n\n\nProceso finalizado. Resumen:")
    unique_questions = {}
    print("Documentos procesados: ", len(documents))
    print("Preguntas totales:", sum([len(qs) for qs in documents]))
    print("Preguntas unicas:", len({q["question"] for qs in documents for q in qs}))
    print("Preguntas+Respuestas unicas: ", 
          len({q["question"]+q["correct_answer"] for qs in documents for q in qs}))
    

print(str(sys.argv))
print("Coincidencias encontradas en:")
print("\n".join(list(found)))
